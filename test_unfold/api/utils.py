"""
Utility functions for dangkiem API
"""
import math


def calculate_distance(lat1, lng1, lat2, lng2):
    """
    Tính khoảng cách giữa 2 điểm sử dụng Haversine formula
    
    Args:
        lat1, lng1: Tọa độ điểm 1 (latitude, longitude)
        lat2, lng2: Tọa độ điểm 2 (latitude, longitude)
    
    Returns:
        float: Khoảng cách (km)
    
    Example:
        >>> distance = calculate_distance(13.776489, 109.223688, 13.780000, 109.230000)
        >>> print(f"{distance:.2f} km")
        0.85 km
    """
    if not all([lat1, lng1, lat2, lng2]):
        return None
    
    # Bán kính Trái Đất (km)
    R = 6371.0
    
    # Chuyển sang radian
    lat1_rad = math.radians(float(lat1))
    lng1_rad = math.radians(float(lng1))
    lat2_rad = math.radians(float(lat2))
    lng2_rad = math.radians(float(lng2))
    
    # Haversine formula
    dlat = lat2_rad - lat1_rad
    dlng = lng2_rad - lng1_rad
    
    a = math.sin(dlat / 2)**2 + math.cos(lat1_rad) * math.cos(lat2_rad) * math.sin(dlng / 2)**2
    c = 2 * math.asin(math.sqrt(a))
    
    distance = R * c
    return distance


def calculate_eta(distance_km, average_speed_kmh=30):
    """
    Tính thời gian ước tính đến nơi (ETA - Estimated Time of Arrival)
    
    Args:
        distance_km: Khoảng cách (km)
        average_speed_kmh: Tốc độ trung bình (km/h) - mặc định 30 km/h cho thành phố
    
    Returns:
        int: Thời gian ước tính (phút)
    
    Example:
        >>> eta = calculate_eta(2.5, average_speed=30)
        >>> print(f"{eta} phút")
        5 phút
    """
    if not distance_km or distance_km <= 0:
        return 0
    
    # Tính thời gian (giờ)
    hours = distance_km / average_speed_kmh
    
    # Chuyển sang phút
    minutes = hours * 60
    
    # Làm tròn lên
    return math.ceil(minutes)


def format_distance(distance_km):
    """
    Format khoảng cách để hiển thị user-friendly
    
    Args:
        distance_km: Khoảng cách (km)
    
    Returns:
        str: Chuỗi hiển thị (VD: "2.5 km", "850 m")
    
    Example:
        >>> format_distance(2.5)
        "2.5 km"
        >>> format_distance(0.35)
        "350 m"
    """
    if not distance_km:
        return "0 km"
    
    if distance_km < 1:
        # Hiển thị mét nếu < 1 km
        meters = int(distance_km * 1000)
        return f"{meters} m"
    else:
        # Hiển thị km với 1 chữ số thập phân
        return f"{distance_km:.1f} km"


def format_eta(minutes):
    """
    Format ETA để hiển thị user-friendly
    
    Args:
        minutes: Thời gian (phút)
    
    Returns:
        str: Chuỗi hiển thị (VD: "5 phút", "1 giờ 30 phút")
    """
    if not minutes or minutes <= 0:
        return "< 1 phút"
    
    if minutes < 60:
        return f"{minutes} phút"
    else:
        hours = minutes // 60
        remaining_minutes = minutes % 60
        if remaining_minutes == 0:
            return f"{hours} giờ"
        else:
            return f"{hours} giờ {remaining_minutes} phút"


# ========================================
# DOCX CONTRACT GENERATION HELPERS
# ========================================

from io import BytesIO
import base64
import glob
import os
import sys
import shutil
import subprocess
import tempfile
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError


def _replace_placeholders_in_paragraph(paragraph, mapping, signature_images=None):
    """Replace placeholders like {{field}} in a paragraph text."""
    if signature_images is None:
        signature_images = {}

    # Handle signature placeholders in paragraphs too
    if '{{customer_signature}}' in paragraph.text:
        paragraph.text = ''
        if 'customer' in signature_images:
            try:
                run = paragraph.add_run()
                run.add_picture(BytesIO(signature_images['customer']), width=Inches(1.5))
            except UnrecognizedImageError as e:
                paragraph.text = '(Chữ ký khách hàng không hợp lệ hoặc bị hỏng)'
            except Exception as e:
                paragraph.text = f"[Chữ ký khách hàng: {type(e).__name__}: {e}]"
        else:
            paragraph.text = '(Chưa có chữ ký khách hàng)'
        return

    if '{{staff_signature}}' in paragraph.text:
        paragraph.text = ''
        if 'staff' in signature_images:
            try:
                run = paragraph.add_run()
                run.add_picture(BytesIO(signature_images['staff']), width=Inches(1.5))
            except Exception as e:
                paragraph.text = f"[Chữ ký nhân viên: {type(e).__name__}: {e}]"
        else:
            paragraph.text = '(Chưa có chữ ký nhân viên)'
        return

    for key, val in mapping.items():
        placeholder = '{{' + key + '}}'
        if placeholder in paragraph.text:
            text = paragraph.text.replace(placeholder, str(val or ''))
            paragraph.text = text


def _replace_placeholders_in_table(table, mapping, signature_images=None):
    """
    Replace placeholders in table cells.

    Args:
        table: docx table object
        mapping: dict of field -> value
        signature_images: dict with 'customer' and 'staff' image bytes
    """
    if signature_images is None:
        signature_images = {}

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Handle signature placeholders first
                if '{{customer_signature}}' in paragraph.text:
                    paragraph.text = ''
                    if 'customer' in signature_images:
                        try:
                            run = paragraph.add_run()
                            run.add_picture(BytesIO(signature_images['customer']), width=Inches(1.5))
                        except UnrecognizedImageError as e:
                            paragraph.text = '(Chữ ký khách hàng không hợp lệ hoặc bị hỏng)'
                        except Exception as e:
                            paragraph.text = f"[Chữ ký khách hàng: {type(e).__name__}: {e}]"
                    else:
                        paragraph.text = '(Chưa có chữ ký khách hàng)'
                    continue

                if '{{staff_signature}}' in paragraph.text:
                    paragraph.text = ''
                    if 'staff' in signature_images:
                        try:
                            run = paragraph.add_run()
                            run.add_picture(BytesIO(signature_images['staff']), width=Inches(1.5))
                        except Exception as e:
                            paragraph.text = f"[Chữ ký nhân viên: {str(e)}]"
                    else:
                        paragraph.text = '(Chưa có)'
                    continue

                # Handle regular text placeholders
                for key, val in mapping.items():
                    placeholder = '{{' + key + '}}'
                    if placeholder in paragraph.text:
                        text = paragraph.text.replace(placeholder, str(val or ''))
                        paragraph.text = text


def _find_soffice_executable() -> str:
    """Find libreoffice/soffice executable on Linux.

    Returns:
        str: path to soffice or libreoffice executable, or None if not found.
    """
    candidates = [
        shutil.which('soffice'),
        shutil.which('libreoffice'),
        '/usr/bin/soffice',
        '/usr/bin/libreoffice',
        '/usr/lib/libreoffice/program/soffice',
    ]

    # Expand some common install paths
    candidates.extend(glob.glob('/opt/libreoffice*/program/soffice'))
    candidates.extend(glob.glob('/snap/bin/soffice'))

    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return candidate
    return None


def convert_docx_bytes_to_pdf(docx_bytes: BytesIO) -> BytesIO:
    """Convert a generated DOCX bytes stream into a PDF bytes stream."""
    if not isinstance(docx_bytes, BytesIO):
        raise TypeError('docx_bytes must be a BytesIO instance')

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_docx_path = os.path.join(tmp_dir, 'contract.docx')
        tmp_pdf_path = os.path.join(tmp_dir, 'contract.pdf')

        with open(tmp_docx_path, 'wb') as tmp_docx_file:
            tmp_docx_file.write(docx_bytes.getvalue())

        if sys.platform.startswith('win'):
            try:
                from docx2pdf import convert as _docx2pdf_convert
            except ImportError as e:
                raise ImportError(
                    'docx2pdf library is required for DOCX to PDF conversion on Windows. Install it with pip install docx2pdf'
                ) from e

            try:
                _docx2pdf_convert(tmp_docx_path, tmp_pdf_path)
            except Exception as e:
                raise Exception(f'Failed to convert DOCX to PDF with docx2pdf: {type(e).__name__}: {e}')

        else:
            # On Linux/Unix, use LibreOffice/soffice directly for conversion.
            soffice_path = _find_soffice_executable()
            if not soffice_path:
                raise EnvironmentError(
                    'Không tìm thấy LibreOffice/soffice trên hệ thống. Cài đặt LibreOffice để chuyển DOCX sang PDF trên Linux.'
                )

            cmd = [soffice_path, '--headless', '--convert-to', 'pdf', '--outdir', tmp_dir, tmp_docx_path]
            try:
                subprocess.run(cmd, capture_output=True, text=True, check=True)
            except subprocess.CalledProcessError as e:
                raise Exception(
                    f'Failed to convert DOCX to PDF with LibreOffice: returncode={e.returncode}; stdout={e.stdout}; stderr={e.stderr}'
                ) from e

        if not os.path.exists(tmp_pdf_path):
            raise FileNotFoundError(f'PDF output not created: {tmp_pdf_path}')

        with open(tmp_pdf_path, 'rb') as tmp_pdf_file:
            pdf_bytes = tmp_pdf_file.read()

    output = BytesIO(pdf_bytes)
    output.seek(0)
    return output


def _decode_base64_image(base64_str: str) -> bytes:
    """Decode base64 string safely, validate image format, fix padding and strip data URI prefix."""
    if not base64_str:
        raise ValueError('Empty base64 string')

    # Remove data URI scheme if provided: data:image/png;base64,...
    if ',' in base64_str:
        base64_str = base64_str.split(',', 1)[1]

    # Remove whitespace/newlines
    base64_str = ''.join(base64_str.split())

    # Fix padding
    padding = len(base64_str) % 4
    if padding:
        base64_str += '=' * (4 - padding)

    # Decode from base64
    try:
        image_bytes = base64.b64decode(base64_str)
    except Exception as e:
        raise ValueError(f"Failed to decode base64: {type(e).__name__}: {str(e)}")
    
    # Validate image data
    if len(image_bytes) < 4:
        raise ValueError(f"Image data too small ({len(image_bytes)} bytes), expected at least 4 bytes")
    
    # Check for valid image format (PNG or JPEG magic numbers)
    if image_bytes[:4] == b'\x89PNG':
        # Valid PNG
        return image_bytes
    elif image_bytes[:3] == b'\xFF\xD8\xFF':
        # Valid JPEG
        return image_bytes
    else:
        # Invalid format - show first bytes in hex
        hex_preview = image_bytes[:20].hex()
        raise ValueError(f"Invalid image format. Expected PNG (89504E47) or JPEG (FFD8FF), got: {hex_preview}")
    
    return image_bytes


def render_contract_docx(template_path: str, data: dict, customer_sig_path: str = None, staff_sig_path: str = None) -> BytesIO:
    """
    Render DOCX contract from template path and data mapping.

    Args:
        template_path: path to .docx template with placeholders {{field}}
        data: dict, e.g. {'customer_name': 'Hoang', 'vehicle_plate': '43A-12345'}
        customer_sig_path: file path to customer signature image (REQUIRED)
        staff_sig_path: file path to staff signature image (OPTIONAL)

    Returns:
        BytesIO containing generated docx file.
    """
    try:
        doc = Document(template_path)
    except Exception as e:
        raise Exception(f"Failed to load template at {template_path}: {type(e).__name__}: {str(e)}")

    # Load signature images from file paths
    signature_images = {}
    
    # Customer signature REQUIRED
    if customer_sig_path:
        if not os.path.exists(customer_sig_path):
            raise FileNotFoundError(f"Customer signature file not found: {customer_sig_path}")

        try:
            with open(customer_sig_path, 'rb') as f:
                customer_bytes = f.read()
        except Exception as e:
            raise Exception(f"Failed to read customer signature file {customer_sig_path}: {type(e).__name__}: {e}")

        try:
            with Image.open(BytesIO(customer_bytes)) as img:
                if img.format not in ('PNG', 'JPEG', 'WEBP'):
                    raise ValueError(f"Customer signature file has unsupported image format: {img.format}")

                if img.format == 'WEBP':
                    with BytesIO() as webp_to_png:
                        img.convert('RGBA').save(webp_to_png, format='PNG')
                        webp_to_png.seek(0)
                        customer_bytes = webp_to_png.read()
        except Exception as e:
            raise Exception(f"Customer signature image invalid: {type(e).__name__}: {e}")

        signature_images['customer'] = customer_bytes
    else:
        raise ValueError('customer_sig_path is required for contract generation')

    # Staff signature OPTIONAL (may not be available yet)
    if staff_sig_path:
        if not os.path.exists(staff_sig_path):
            print(f"Warning: Staff signature file not found: {staff_sig_path}")
        else:
            try:
                with open(staff_sig_path, 'rb') as f:
                    staff_bytes = f.read()
            except Exception as e:
                print(f"Warning: Failed to read staff signature file from {staff_sig_path}: {type(e).__name__}: {e}")
            else:
                try:
                    with Image.open(BytesIO(staff_bytes)) as img:
                        if img.format not in ('PNG', 'JPEG', 'WEBP'):
                            print(f"Warning: Staff signature has unsupported image format: {img.format}")
                        else:
                            if img.format == 'WEBP':
                                with BytesIO() as webp_to_png:
                                    img.convert('RGBA').save(webp_to_png, format='PNG')
                                    webp_to_png.seek(0)
                                    signature_images['staff'] = webp_to_png.read()
                            else:
                                signature_images['staff'] = staff_bytes
                except Exception as e:
                    print(f"Warning: Staff signature invalid image: {e}")

    # Replace in paragraphs and tables
    try:
        for paragraph in doc.paragraphs:
            _replace_placeholders_in_paragraph(paragraph, data, signature_images)

        for table in doc.tables:
            _replace_placeholders_in_table(table, data, signature_images)
    except Exception as e:
        raise Exception(f"Failed to replace placeholders in document: {type(e).__name__}: {str(e)}")

    try:
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    except Exception as e:
        raise Exception(f"Failed to save document to BytesIO: {type(e).__name__}: {str(e)}")

