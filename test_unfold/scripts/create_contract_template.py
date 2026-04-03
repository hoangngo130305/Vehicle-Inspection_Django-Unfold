from docx import Document
import os

# Inputs
base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # test_unfold
templates_dir = os.path.join(base_dir, 'templates')
if not os.path.exists(templates_dir):
    os.makedirs(templates_dir)

output_path = os.path.join(templates_dir, 'contract_template.docx')

# Build document
doc = Document()

# Header + title
h1 = doc.add_paragraph()
h1.alignment = 1  # center
r1 = h1.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n')
r1.bold = True
r1.font.size = 160000

h2 = doc.add_paragraph()
h2.alignment = 1
r2 = h2.add_run('ĐỘC LẬP - TỰ DO - HẠNH PHÚC')
r2.bold = True
r2.font.size = 140000

title = doc.add_paragraph()
title.alignment = 1
rt = title.add_run('HỢP ĐỒNG ỦY QUYỀN\n')
rt.bold = True
rt.font.size = 220000

subtitle = doc.add_paragraph()
subtitle.alignment = 1
rs = subtitle.add_run('(V/v: Nhận xe ô tô và thực hiện thủ tục đăng kiểm)')
rs.italic = True

#Body
p = doc.add_paragraph('Hôm nay, ngày {{today_date}}, tại {{station_address}}, chúng tôi gồm có:')
p.style = 'Body Text'

p = doc.add_paragraph()
r = p.add_run('1. BÊN ỦY QUYỀN (BÊN A – CHỦ XE)')
r.bold = True

doc.add_paragraph('Họ và tên: {{customer_name}}')

doc.add_paragraph('Ngày sinh: {{customer_date_of_birth}}')

doc.add_paragraph('Số CCCD/CMND/Hộ chiếu: {{customer_id_number}} cấp ngày {{customer_id_issued_date}} nơi cấp {{customer_id_issued_place}}')

doc.add_paragraph('Địa chỉ thường trú: {{customer_address}}')

doc.add_paragraph('Số điện thoại: {{customer_phone}}')

doc.add_paragraph('Là chủ sở hữu hợp pháp của xe ô tô có thông tin sau:')

doc.add_paragraph('Nhãn hiệu xe: {{vehicle_manufacturer}}')

doc.add_paragraph('Biển số: {{vehicle_plate}}')

doc.add_paragraph('Số khung: {{vehicle_chassis_number}}')

doc.add_paragraph('Số máy: {{vehicle_engine_number}}')

p = doc.add_paragraph()
r = p.add_run('2. BÊN ĐƯỢC ỦY QUYỀN (BÊN B – ĐƠN VỊ DỊCH VỤ)')
r.bold = True

p = doc.add_paragraph()
r = p.add_run('Tên đơn vị: TRUNG TÂM HỖ TRỢ DỊCH VỤ ĐĂNG KIỂM VIỆT DKV 50S')
r.bold = True

doc.add_paragraph('Mã số doanh nghiệp: 0316969591 - 00005')

doc.add_paragraph('Địa chỉ trụ sở: 26B Đường 34 - Phường Thủ Đức – TP.HCM')

doc.add_paragraph('Đại diện theo pháp luật: Đặng Hồng Nam - Chức vụ: Giám Đốc')

doc.add_paragraph('Số điện thoại: 0944484444')

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('ĐIỀU 1. NỘI DUNG VÀ PHẠM VI ỦY QUYỀN')
r.bold = True

doc.add_paragraph('Bên A đồng ý ủy quyền cho Bên B thực hiện các công việc sau:')

items_1 = [
    'Nhận xe ô tô nêu trên tại địa chỉ do Bên A chỉ định.',
    'Thay mặt Bên A điều khiển xe chỉ nhằm mục đích đưa xe đi đăng kiểm và đưa xe trở lại.',
    'Thực hiện các thủ tục đăng kiểm xe cơ giới theo quy định pháp luật.',
    'Nộp các khoản phí, lệ phí đăng kiểm (nếu có) theo thỏa thuận giữa hai bên.',
    'Nhận lại Giấy chứng nhận kiểm định và tem kiểm định để bàn giao cho Bên A.'
]
for item in items_1:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

doc.add_paragraph('👉 Bên B không được sử dụng xe vào bất kỳ mục đích nào khác ngoài các nội dung nêu trên.')

p = doc.add_paragraph()
r = p.add_run('ĐIỀU 2. THỜI HẠN ỦY QUYỀN')
r.bold = True

doc.add_paragraph('Thời hạn ủy quyền: từ ngày {{authorization_start_date}} đến hết ngày {{authorization_end_date}}')

doc.add_paragraph('Văn bản ủy quyền tự động chấm dứt hiệu lực sau khi Bên B hoàn thành việc bàn giao xe và giấy tờ liên quan cho Bên A.')

p = doc.add_paragraph()
r = p.add_run('ĐIỀU 3. CAM KẾT CỦA BÊN B (ĐƠN VỊ DỊCH VỤ)')
r.bold = True

items_3 = [
    'Thực hiện đúng phạm vi ủy quyền, tuân thủ luật giao thông đường bộ.',
    'Chịu trách nhiệm đối với các vi phạm giao thông phát sinh do lỗi của Bên B trong thời gian nhận và điều khiển xe.',
    'Bồi thường thiệt hại nếu xảy ra mất mát, hư hỏng xe do lỗi của Bên B.',
    'Không giao xe cho bên thứ ba khi chưa có sự đồng ý của Bên A (trừ trường hợp nhân sự của công ty thực hiện theo phân công nội bộ).'
]
for item in items_3:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)


p = doc.add_paragraph()
r = p.add_run('ĐIỀU 4. CAM KẾT CỦA BÊN A (CHỦ XE)')
r.bold = True

items_4 = [
    'Cam kết xe đủ điều kiện lưu hành, không tranh chấp, không bị cầm cố, thế chấp trái pháp luật.',
    'Cung cấp đầy đủ, trung thực giấy tờ liên quan đến xe (đăng ký xe, bảo hiểm, giấy tờ khác nếu có).',
    'Chịu trách nhiệm đối với các lỗi kỹ thuật, tình trạng xe không đạt đăng kiểm không phát sinh từ quá trình vận chuyển của Bên B.',
    'Thanh toán đầy đủ chi phí dịch vụ theo thỏa thuận.'
]
for item in items_4:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

p = doc.add_paragraph()
r = p.add_run('ĐIỀU 5. GIỚI HẠN TRÁCH NHIỆM')
r.bold = True

doc.add_paragraph('Bên B không chịu trách nhiệm đối với:')

items_5 = [
    'Các hư hỏng, sự cố do lỗi kỹ thuật có sẵn của xe.',
    'Việc xe không đạt đăng kiểm do nguyên nhân khách quan hoặc tình trạng xe.',
    'Sự kiện bất khả kháng (tai nạn không do lỗi, thiên tai, sự cố giao thông ngoài tầm kiểm soát).'
]
for item in items_5:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

p = doc.add_paragraph()
r = p.add_run('ĐIỀU 6. HIỆU LỰC')
r.bold = True

doc.add_paragraph('Văn bản này được lập thành 02 bản, mỗi bên giữ 01 bản, có giá trị pháp lý như nhau.')

doc.add_paragraph('Hai bên đã đọc, hiểu rõ quyền và nghĩa vụ của mình và tự nguyện ký tên dưới đây.')

# Signature section
sig_table = doc.add_table(rows=2, cols=2)
sig_table.autofit = False

cell_a = sig_table.rows[0].cells[0]
cell_a_para = cell_a.paragraphs[0]
r = cell_a_para.add_run('ĐẠI DIỆN BÊN A\n(Ký, ghi rõ họ tên)')
r.bold = True

cell_b = sig_table.rows[0].cells[1]
cell_b_para = cell_b.paragraphs[0]
r = cell_b_para.add_run('ĐẠI DIỆN BÊN B\n(Ký, ghi rõ họ tên, đóng dấu công ty)')
r.bold = True

sig_table.rows[1].cells[0].text = ''
sig_table.rows[1].cells[1].text = 'ĐẶNG HỒNG NAM'

doc.save(output_path)
print('Template contract_template.docx đã được tạo tại', output_path)
