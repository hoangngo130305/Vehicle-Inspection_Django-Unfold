from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os


def add_center_bold_paragraph(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def add_label_value(doc, label, placeholder):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(placeholder)


def build_handover_template(output_path):
    doc = Document()

    # Header
    add_center_bold_paragraph(doc, "CONG HOA XA HOI CHU NGHIA VIET NAM", 12)
    add_center_bold_paragraph(doc, "Doc lap - Tu do - Hanh phuc", 12)
    doc.add_paragraph()

    add_center_bold_paragraph(doc, "BIEN BAN BAN GIAO", 14)
    doc.add_paragraph()

    # Intro
    doc.add_paragraph(
        "Can cu Hop dong so {{contract_number}} ngay {{contract_day}}/{{contract_month}}/{{contract_year}}, "
        "duoc ky ket giua {{party_a_name}} va {{service_company_name}}."
    )
    doc.add_paragraph(
        "Hom nay, ngay {{handover_day}} thang {{handover_month}} nam {{handover_year}}, "
        "tai {{handover_location}}, chung toi gom co:"
    )

    # Party A
    p = doc.add_paragraph()
    p.add_run("DAI DIEN BEN A (CHU XE)").bold = True
    add_label_value(doc, "Ho va ten", "{{customer_name}}")
    add_label_value(doc, "Ngay sinh", "{{customer_date_of_birth}}")
    add_label_value(doc, "So CCCD/CMND/Ho chieu", "{{customer_id_number}}")
    add_label_value(doc, "Ngay cap", "{{customer_id_issued_date}}")
    add_label_value(doc, "Noi cap", "{{customer_id_issued_place}}")
    add_label_value(doc, "Dia chi thuong tru", "{{customer_address}}")
    add_label_value(doc, "So dien thoai", "{{customer_phone}}")

    doc.add_paragraph("La chu so huu hop phap cua xe o to co thong tin sau:")
    add_label_value(doc, "Nhan hieu xe", "{{vehicle_brand}}")
    add_label_value(doc, "Bien so", "{{vehicle_plate}}")
    add_label_value(doc, "So khung", "{{vehicle_chassis_number}}")
    add_label_value(doc, "So may", "{{vehicle_engine_number}}")

    # Party B
    p = doc.add_paragraph()
    p.add_run("DAI DIEN BEN B (DON VI DICH VU)").bold = True
    add_label_value(doc, "Ten don vi", "{{service_company_name}}")
    add_label_value(doc, "Ma so doanh nghiep", "{{service_company_tax_code}}")
    add_label_value(doc, "Dia chi tru so", "{{service_company_address}}")
    add_label_value(doc, "Dai dien cong ty", "{{staff_name}}")
    add_label_value(doc, "Chuc vu", "{{staff_role}}")
    add_label_value(doc, "So dien thoai", "{{staff_phone}}")

    doc.add_paragraph(
        "Hai ben tien hanh giao nhan xe o to bien so {{vehicle_plate}} va hien trang thuc te cua xe chi tiet nhu sau:"
    )

    # Checklist table 9 items
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["STT", "Ngoai quan cua xe", "Khong dat", "Dat", "So luong", "Ghi chu"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    checklist_rows = [
        (1, "Xe co bi tray xuot", "{{scratches_not_passed}}", "{{scratches_passed}}", "{{scratches_quantity}}", "{{scratches_note}}"),
        (2, "Lop", "{{tires_not_passed}}", "{{tires_passed}}", "{{tires_quantity}}", "{{tires_note}}"),
        (3, "Phanh", "{{brakes_not_passed}}", "{{brakes_passed}}", "{{brakes_quantity}}", "{{brakes_note}}"),
        (4, "Binh", "{{battery_not_passed}}", "{{battery_passed}}", "{{battery_quantity}}", "{{battery_note}}"),
        (5, "Catvet", "{{carpet_not_passed}}", "{{carpet_passed}}", "{{carpet_quantity}}", "{{carpet_note}}"),
        (6, "Dang kiem", "{{inspection_not_passed}}", "{{inspection_passed}}", "{{inspection_quantity}}", "{{inspection_note}}"),
        (7, "Bao hiem", "{{insurance_not_passed}}", "{{insurance_passed}}", "{{insurance_quantity}}", "{{insurance_note}}"),
        (8, "Khoi", "{{smoke_not_passed}}", "{{smoke_passed}}", "{{smoke_quantity}}", "{{smoke_note}}"),
        (9, "Den", "{{lights_not_passed}}", "{{lights_passed}}", "{{lights_quantity}}", "{{lights_note}}"),
    ]

    for row in checklist_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)

    doc.add_paragraph()
    doc.add_paragraph(
        "Thiet bi xe da qua su dung, cac hinh anh se duoc ben B chup lai va gui qua zalo so {{zalo_phone_confirm}} "
        "de ben A xac nhan, dung theo bien ban va dong thoi lam co so de hai ben giao nhan xe."
    )
    doc.add_paragraph(
        "Ben A va Ben B da cung kiem tra xe va xac nhan ngoai quan thuc te day du nhu tren khi giao xe ben B."
    )
    doc.add_paragraph(
        "Bien ban nay duoc lap thanh 02 ban co gia tri phap ly nhu nhau, Ben A giu 01 ban, Ben B giu 01 ban."
    )

    # Signature block 1
    sig1 = doc.add_table(rows=1, cols=2)
    sig1.style = "Table Grid"
    sig1.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig1.rows[0].cells[0].text = "DAI DIEN BEN A\n(Ky, ghi ro ho ten)\n\n{{customer_signature}}\n{{customer_name}}"
    sig1.rows[0].cells[1].text = "DAI DIEN BEN B\n(Ky, ghi ro ho ten)\n\n{{staff_signature}}\n{{staff_name}}"

    doc.add_paragraph()
    doc.add_paragraph(
        "XAC NHAN BEN B DA DANG KIEM XONG, BAN GIAO TEM DANG KIEM VA GIAY TO CHO BEN A. "
        "BEN A DA NHAN BAN GIAO XE TU BEN B VAO LUC {{return_hour}} gio {{return_minute}} phut "
        "ngay {{return_day}} thang {{return_month}} nam {{return_year}}. "
        "Dong y ky bien ban thanh ly hop dong so {{contract_number}} giua {{party_a_name}} "
        "va {{service_company_name}}."
    )

    # Signature block 2
    sig2 = doc.add_table(rows=1, cols=2)
    sig2.style = "Table Grid"
    sig2.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig2.rows[0].cells[0].text = "DAI DIEN BEN A\n(Ky, ghi ro ho ten)\n\n{{customer_signature}}\n{{customer_name}}"
    sig2.rows[0].cells[1].text = "DAI DIEN BEN B\n(Ky, ghi ro ho ten)\n\n{{staff_signature}}\n{{staff_name}}"

    doc.save(output_path)


def main():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    templates_dir = os.path.join(base_dir, "templates")
    os.makedirs(templates_dir, exist_ok=True)

    output_path = os.path.join(templates_dir, "handover_template.docx")
    build_handover_template(output_path)
    print(f"Template handover_template.docx da duoc tao tai: {output_path}")


if __name__ == "__main__":
    main()